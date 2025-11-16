from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path

SECTIONS = [
    ("ภาพรวม (Overview)", "Appium คือ Automation Framework สำหรับ Mobile ใช้มาตรฐาน WebDriver"),
    ("Prerequisites", "JDK, Android SDK, Node.js, Appium 2.x, Python, VS Code, Emulator/Device"),
    ("การติดตั้งเครื่องมือ", "สรุปคำสั่งติดตั้ง JDK / Node / Appium / pip libs"),
    ("การตั้งค่า Environment Variables", "JAVA_HOME ANDROID_HOME PATH"),
    ("การเตรียมอุปกรณ์ (Device)", "adb devices, emulator -avd"),
    ("การเปิด Appium Server", "appium --port 4723"),
    ("การใช้ Appium Inspector", "กำหนด capabilities และค้นหา locator"),
    ("โครงสร้างโปรเจกต์", "tests pages utils config reports screenshots"),
    ("Desired Capabilities", "platformName automationName appPackage appActivity noReset unicodeKeyboard"),
    ("สคริปต์พื้นฐาน (Demo)", "basic_demo.py"),
    ("Pytest + Page Object Pattern", "แยก element locator จาก logic"),
    ("Locator Strategies", "resource-id > accessibility id > UIAutomator > xpath"),
    ("Waiting Strategies", "Explicit wait"),
    ("Unicode / ภาษาไทย", "unicodeKeyboard resetKeyboard"),
    ("Screenshot & Logs", "save_screenshot page_source"),
    ("Permission Handling", "autoGrantPermissions"),
    ("Performance Tips", "ปิด animation ใช้ noReset parallel"),
    ("Utilities", "waits.py logger.py"),
    ("Troubleshooting", "ปัญหาพบบ่อย + adb")
]

def build():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'TH Sarabun New'
    normal.font.size = Pt(14)
    normal.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')

    doc.add_heading('คู่มือการติดตั้งและใช้งาน Appium Android (Python + VS Code)', 0)
    doc.add_paragraph('เวอร์ชัน 1.0\nผู้จัดทำ: (ใส่ชื่อ)\nวันที่: (ใส่วันที่)')
    doc.add_page_break()

    doc.add_heading('สารบัญ (Insert > Table of Contents ใน Word)', level=1)
    doc.add_page_break()

    for title, body in SECTIONS:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)

    doc.add_page_break()
    doc.add_heading('ภาคผนวก A – โค้ด', level=1)
    doc.add_paragraph('basic_demo.py test_sample_login.py login_page.py home_page.py capabilities.json')

    out = Path('Appium_Full_Guide_generated.docx')
    doc.save(out)
    print(f'Saved {out.resolve()}')

if __name__ == '__main__':
    build()
